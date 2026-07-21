from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "GUIA_USUARIO_HERMES.md"
OUTPUT = ROOT / "GUIA_USUARIO_HERMES.docx"
BLUE = RGBColor(35, 131, 196)
DARK = RGBColor(31, 41, 55)
MUTED = RGBColor(90, 101, 116)


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    run.font.size = Pt(8)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run._r.addnext(field)


def add_rich_text(paragraph, text: str) -> None:
    for part in re.split(r"(\*\*.+?\*\*)", text):
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        run = paragraph.add_run(part[2:-2] if bold else part)
        run.bold = bold


def configure(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    for style_name, size, color in (
        ("Title", 32, DARK), ("Subtitle", 16, MUTED),
        ("Heading 1", 20, BLUE), ("Heading 2", 15, DARK), ("Heading 3", 12, BLUE),
    ):
        style = document.styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def add_cover(document: Document) -> None:
    document.add_paragraph()
    document.add_paragraph()
    bar = document.add_table(rows=1, cols=1)
    bar.autofit = False
    bar.columns[0].width = Cm(16)
    bar.rows[0].height = Cm(0.22)
    set_cell_shading(bar.cell(0, 0), "2383C4")

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("HERMES")
    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.add_run("Manual do Usuário")
    description = document.add_paragraph()
    description.paragraph_format.space_before = Pt(14)
    run = description.add_run("Painel de Automação de Planilhas")
    run.font.size = Pt(16)
    run.font.color.rgb = BLUE
    version = document.add_paragraph()
    version.paragraph_format.space_before = Pt(180)
    version.add_run("Versão do guia: julho de 2026").font.color.rgb = MUTED
    document.add_page_break()


def add_header_footer(document: Document) -> None:
    section = document.sections[0]
    header = section.header.paragraphs[0]
    header.text = "HERMES  |  Manual do Usuário"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = MUTED
    add_page_number(section.footer.paragraphs[0])


def build() -> None:
    document = Document()
    configure(document)
    add_cover(document)
    add_header_footer(document)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    start = next(index for index, line in enumerate(lines) if line.startswith("## 1."))
    for raw in lines[start:]:
        line = raw.strip()
        if not line or line == "---":
            continue
        if line.startswith("### "):
            document.add_heading(line[4:], level=2)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=1)
        elif re.match(r"^\d+\. ", line):
            paragraph = document.add_paragraph(style="List Number")
            add_rich_text(paragraph, re.sub(r"^\d+\. ", "", line))
        elif line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_rich_text(paragraph, line[2:])
        elif line.startswith("> "):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.7)
            paragraph.paragraph_format.space_before = Pt(5)
            paragraph.paragraph_format.space_after = Pt(8)
            add_rich_text(paragraph, line[2:])
            paragraph.runs[0].font.color.rgb = MUTED
        else:
            paragraph = document.add_paragraph()
            add_rich_text(paragraph, line.rstrip("  "))

    document.save(OUTPUT)


if __name__ == "__main__":
    build()
